
from __future__ import annotations

from pathlib import Path
from typing import Callable, Optional

import fitz  # PyMuPDF
from docx import Document

from src.core.unlocker import open_document
from src.core.ocr import ocr_page
from src.utils.logger import get_logger
from src.utils.file_ops import ensure_parent_dirs
from src.config import WORD_DOCS_DIR

log = get_logger(__name__)

ProgressCB = Optional[Callable[[str], None]]


def _emit(cb: ProgressCB, message: str) -> None:
    log.info(message)
    if cb:
        try:
            cb(message)
        except Exception:
            log.exception("Progress callback failed.")


def convert_pdf_to_docx(
    pdf_path: str | Path,
    username: str | None = None,
    password: str | None = None,
    progress_cb: ProgressCB = None,
    ocr_lang: str | None = None,
    min_text_threshold: int = 30,
) -> Path:
    """
    Convert a single PDF to a .docx file saved under WORD_DOCS_DIR.

    Args:
        pdf_path: path-like to the input PDF.
        username/password: credentials for encrypted PDFs.
        progress_cb: optional callable(str) to relay progress to GUI.
        ocr_lang: optional Tesseract language (e.g., 'eng+spa').
        min_text_threshold: if extracted text length < this, OCR fallback.

    Returns:
        Path to the generated .docx file.
    """
    pdf_path = Path(pdf_path).expanduser().resolve()
    if not pdf_path.exists():
        raise FileNotFoundError(pdf_path)

    _emit(progress_cb, f"Opening PDF: {pdf_path.name}")
    doc = open_document(pdf_path, username=username, password=password)

    output_path = WORD_DOCS_DIR / f"{pdf_path.stem}.docx"
    ensure_parent_dirs([output_path])

    _emit(progress_cb, "Beginning text extraction...")
    word_doc = Document()
    word_doc.core_properties.title = pdf_path.stem

    for page_index in range(doc.page_count):
        page = doc.load_page(page_index)
        _emit(progress_cb, f"Processing page {page_index + 1}/{doc.page_count}")

        text = page.get_text().strip()

        if len(text) < min_text_threshold:
            _emit(progress_cb, f"Page {page_index + 1}: insufficient text ({len(text)} chars); running OCR.")
            try:
                text = ocr_page(page, lang=ocr_lang)
            except Exception as exc:
                log.exception("OCR failed on page %s: %s", page_index + 1, exc)
        else:
            log.debug("Page %s extracted via direct text.", page_index + 1)

        if text:
            for line in text.splitlines():
                word_doc.add_paragraph(line)
        else:
            word_doc.add_paragraph("[Blank page]")

        if page_index < doc.page_count - 1:
            word_doc.add_page_break()

    doc.close()

    word_doc.save(output_path)
    _emit(progress_cb, f"Saved Word document: {output_path.name}")
    return output_path
